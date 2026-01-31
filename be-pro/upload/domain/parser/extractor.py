from __future__ import annotations

import os
import threading
from concurrent.futures import as_completed
from concurrent.futures import ThreadPoolExecutor
from io import BytesIO
from typing import List
from typing import Tuple

import cv2
import docx
import numpy as np
import openpyxl
import pdfplumber
import pytesseract
from fastapi import UploadFile
from pdf2image import convert_from_bytes
from pdfplumber.utils import extract_text
from pdfplumber.utils import get_bbox_overlap
from pdfplumber.utils import obj_to_bbox
from PIL import Image

from .base import FileType


class ExtractorService:
    """Service for extracting raw text from various file formats."""

    def extract(self, file: UploadFile) -> str:
        """Extract raw text from the given file based on its format."""
        file_type = self.get_file_type(file)

        extraction_methods = {
            FileType.PDF: self.extract_pdf,
            FileType.DOCX: self.extract_docx,
            FileType.XLSX: self.extract_xlsx,
            FileType.IMAGE: self.extract_image,
        }

        if file_type in extraction_methods:
            return extraction_methods[file_type](file)
        else:
            raise ValueError(f'Unsupported file type: {file_type}')

    def get_file_type(self, file: UploadFile) -> FileType:
        """Determine file type based on filename and content type."""
        if not file.filename:
            return FileType.UNKNOWN

        _, ext = os.path.splitext(file.filename.lower())

        type_mapping = {
            '.pdf': FileType.PDF,
            '.docx': FileType.DOCX,
            '.xlsx': FileType.XLSX,
            '.xls': FileType.XLSX,
            '.jpg': FileType.IMAGE,
            '.jpeg': FileType.IMAGE,
            '.png': FileType.IMAGE,
            '.gif': FileType.IMAGE,
            '.bmp': FileType.IMAGE,
            '.tiff': FileType.IMAGE,
        }

        return type_mapping.get(ext, FileType.UNKNOWN)

    def extract_docx(self, file: UploadFile) -> str:
        """Extract from a DOCX file."""
        try:
            file.file.seek(0)
            document = docx.Document(file.file)
            content = []

            # Extract paragraphs and tables in document order
            for element in document.element.body:
                if element.tag.endswith('p'):  # Paragraph
                    # Find corresponding paragraph object
                    for para in document.paragraphs:
                        if para._element == element:
                            text = para.text.strip()
                            if text:
                                content.append(text)
                            break
                elif element.tag.endswith('tbl'):  # Table
                    # Find corresponding table object
                    for table in document.tables:
                        if table._element == element:
                            md_table = self.__docx_table_to_markdown(table)
                            if md_table:
                                content.append(md_table)
                            break

            # Handle images (stub)
            for rel in document.part.rels.values():
                if 'image' in rel.target_ref:
                    image_content = self.__ocr_image(np.array(Image.open(BytesIO(rel.target_part.blob))))
                    content.append(image_content)

            return '\n\n'.join(content)
        except Exception as e:
            raise ValueError(f'Failed to extract text from DOCX file: {str(e)}')

    def __docx_table_to_markdown(self, table) -> str:
        """Convert DOCX table to Markdown format."""
        if not table.rows:
            return ''

        rows_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                # Clean cell text and handle None values
                cell_text = (cell.text or '').strip().replace('\n', ' ')
                row_data.append(cell_text)
            rows_data.append(row_data)

        if not rows_data:
            return ''

        # Create markdown table
        header = '| ' + ' | '.join(rows_data[0]) + ' |'
        separator = '| ' + ' | '.join(['---'] * len(rows_data[0])) + ' |'
        body_rows = ['| ' + ' | '.join(row) + ' |' for row in rows_data[1:]]

        return '\n'.join([header, separator] + body_rows)

    def __convert_pdf_page_to_image(self, file_byte: bytes, page_number: int) -> np.ndarray:
        """Convert PDF file bytes to a single PIL Image object."""
        try:
            image = convert_from_bytes(pdf_file=file_byte, first_page=page_number, last_page=page_number, dpi=300, fmt='jpg')[0]
            return np.array(image)
        except Exception as e:
            raise ValueError(f'Failed to convert PDF to images: {str(e)}')

    def __process_single_page(self, page_data: Tuple[int, bytes, UploadFile]) -> Tuple[int, str]:
        page_index, file_byte, file = page_data
        try:
            with pdfplumber.open(BytesIO(file_byte)) as pdf:
                if page_index >= len(pdf.pages):
                    raise IndexError('Page index out of range')

                page = pdf.pages[page_index]
                filtered_page = page
                chars = filtered_page.chars

                for table in page.find_tables():
                    table_chars = page.crop(table.bbox).chars
                    if not table_chars:
                        continue
                    first_table_char = table_chars[0]
                    filtered_page = filtered_page.filter(
                        lambda obj: get_bbox_overlap(obj_to_bbox(obj), table.bbox) is None,
                    )
                    chars = filtered_page.chars
                    markdown = self.__table_to_markdown(table.extract())
                    chars.append(first_table_char | {'text': markdown})

                page_text = extract_text(chars, layout=True)

                if not page_text:
                    page_text = ''
                    image = self.__convert_pdf_page_to_image(file_byte, page_index + 1)
                    ocr_text = self.__ocr_image(image)
                    if ocr_text:
                        page_text += ocr_text + '\n'

            return (page_index, page_text)

        except Exception as e:
            print(f'Lỗi xử lý trang {page_index + 1}: {str(e)}')
            return (page_index, '')

    def extract_pdf(self, file: UploadFile) -> str:
        """Extract from a PDF file."""
        try:
            file.file.seek(0)
            file_byte = file.file.read()

            with pdfplumber.open(BytesIO(file_byte)) as pdf:
                total_pages = len(pdf.pages)

            page_data_list = [(i, file_byte, file) for i in range(total_pages)]

            results = {}
            with ThreadPoolExecutor(max_workers=4) as executor:
                future_to_index = {
                    executor.submit(self.__process_single_page, page_data): page_data[0]
                    for page_data in page_data_list
                }

                for future in as_completed(future_to_index):
                    try:
                        page_index, page_text = future.result()
                        results[page_index] = page_text.strip()
                    except Exception as e:
                        page_index = future_to_index[future]
                        print(f'Lỗi xử lý trang {page_index + 1}: {str(e)}')
                        results[page_index] = ''

            content = []
            for i in sorted(results.keys()):
                content.append(results[i])

            return '\n'.join(content)

        except Exception as e:
            raise ValueError(f'Failed to extract text from PDF file: {str(e)}')

    def __table_to_markdown(self, table: List[List[str]]) -> str:
        """Convert table to Markdown format."""
        if not table or not table[0]:
            return ''

        # Clean table data
        cleaned_table = []
        for row in table:
            cleaned_row = []
            for cell in row:
                if cell is None:
                    cleaned_row.append('')
                else:
                    # Clean cell content
                    cleaned_cell = str(cell).strip().replace('\n', ' ')
                    cleaned_row.append(cleaned_cell)
            cleaned_table.append(cleaned_row)

        # Ensure all rows have the same number of columns
        max_cols = max(len(row) for row in cleaned_table)
        for row in cleaned_table:
            while len(row) < max_cols:
                row.append('')

        # Create markdown table
        header = '| ' + ' | '.join(cleaned_table[0]) + ' |'
        separator = '| ' + ' | '.join(['---'] * max_cols) + ' |'
        body_rows = ['| ' + ' | '.join(row) + ' |' for row in cleaned_table[1:]]

        return '\n'.join([header, separator] + body_rows)

    def __ocr_image(self, image: np.ndarray) -> str:
        try:
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

            # Làm mờ để giảm nhiễu
            blurred = cv2.GaussianBlur(gray, (5, 5), 0)
            # Áp dụng adaptive threshold để tạo ảnh nhị phân
            thresh = cv2.adaptiveThreshold(
                blurred, 255,
                cv2.ADAPTIVE_THRESH_MEAN_C,
                cv2.THRESH_BINARY_INV,
                15, 4,
            )

            # Tìm các đường thẳng dọc và ngang
            kernel_h = cv2.getStructuringElement(cv2.MORPH_RECT, (30, 1))
            kernel_v = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 30))

            detect_h = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, kernel_h, iterations=2)
            detect_v = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, kernel_v, iterations=2)

            # Kết hợp cả hai để tạo mask bảng
            table_mask = cv2.add(detect_h, detect_v)

            # Tìm contours từ mask bảng
            contours, _ = cv2.findContours(table_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

            # Cắt bảng và phủ trắng + placeholder
            table_images: list[np.ndarray] = []
            for i, cnt in enumerate(contours):
                x, y, w, h = cv2.boundingRect(cnt)
                if w > 100 and h > 100:
                    table = image[y:y + h, x:x + w].copy()
                    table_images.insert(0, table)

                    # Tạo placeholder
                    cv2.rectangle(image, (x, y), (x + w, y + h), (255, 255, 255), thickness=-1)

                    text = 'table here'
                    font = cv2.FONT_HERSHEY_SIMPLEX
                    font_scale = 1
                    thickness = 2
                    text_size = cv2.getTextSize(text, font, font_scale, thickness)[0]
                    text_x = x + (w - text_size[0]) // 2
                    text_y = y + (h + text_size[1]) // 2
                    cv2.putText(image, text, (text_x, text_y), font, font_scale, (0, 0, 0), thickness)

            custom_config = r'--oem 3 --psm 6 -l vie'
            page_content = pytesseract.image_to_string(image, config=custom_config)

            # If Vietnamese fails, try English only
            if not page_content:
                custom_config = r'--oem 3 --psm 6 -l eng'
                page_content = pytesseract.image_to_string(image, config=custom_config)

            table_contents = []
            for table_image in table_images:
                cells = self.__extract_cells_from_table(table_image)
                table_data = []
                for row in cells:
                    row_data = []
                    for cell in row:
                        cell_content = pytesseract.image_to_string(cell, config=custom_config)
                        # Cho phép xuống dòng trong ô:
                        cell_content = cell_content.replace('\n', '<br>')
                        row_data.append(cell_content)
                    table_data.append(row_data)

                markdown = self.__table_to_markdown(table_data)
                table_contents.append(markdown)

            for content in table_contents:
                page_content = page_content.replace('table here', content, 1)

            return page_content

        except Exception as e:
            raise ValueError(f'Failed to extract text from image: {str(e)}')

    def __extract_cells_from_table(self, table_image: np.ndarray):
        gray = cv2.cvtColor(table_image, cv2.COLOR_BGR2GRAY)
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        contrast = clahe.apply(gray)
        blur = cv2.GaussianBlur(contrast, (3, 3), 0)

        # Phát hiện đường kẻ bằng adaptive threshold
        thresh = cv2.adaptiveThreshold(
            blur, 255, cv2.ADAPTIVE_THRESH_MEAN_C,
            cv2.THRESH_BINARY_INV, 15, 4,
        )

        # Dilation để nối border kép thành một khối
        dilated = cv2.dilate(thresh, np.ones((3, 3), np.uint8), iterations=1)

        # Phát hiện các đường kẻ dọc và ngang
        scale = 20
        h_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (table_image.shape[1] // scale, 1))
        v_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, table_image.shape[0] // scale))

        detect_h = cv2.morphologyEx(dilated, cv2.MORPH_OPEN, h_kernel, iterations=2)
        detect_v = cv2.morphologyEx(dilated, cv2.MORPH_OPEN, v_kernel, iterations=2)

        # Kết hợp để tạo mặt nạ bảng
        grid_mask = cv2.add(detect_h, detect_v)
        merged_mask = cv2.dilate(grid_mask, np.ones((3, 3), np.uint8), iterations=1)

        # Tìm contour của ô
        contours, _ = cv2.findContours(merged_mask, cv2.RETR_TREE, cv2.CHAIN_APPROX_SIMPLE)

        cell_coords = []
        for cnt in contours:
            x, y, w, h = cv2.boundingRect(cnt)
            if w > 20 and h > 20:  # lọc nhiễu nhỏ
                cell_coords.append((x, y, w, h))

        img_h, img_w = table_image.shape[:2]

        for cnt in contours:
            x, y, w, h = cv2.boundingRect(cnt)
            area = w * h
            if area > 0.9 * img_h * img_w:
                cell_coords.remove((x, y, w, h))

        # Sắp xếp theo y trước, x sau
        cell_coords = sorted(cell_coords, key=lambda b: (b[1], b[0]))

        # Gom các ô thành từng hàng
        rows = []
        current_row: list[tuple[int, int, int, int]] = []
        last_y = -100
        tolerance_y = 20

        for bbox in cell_coords:
            x, y, w, h = bbox
            if abs(y - last_y) > tolerance_y:
                if current_row:
                    rows.append(current_row)
                current_row = [bbox]
                last_y = y
            else:
                current_row.append(bbox)
        if current_row:
            rows.append(current_row)

        # Sắp xếp các ô trong từng hàng theo x
        cells = []
        for row in rows:
            row_sorted = sorted(row, key=lambda b: b[0])
            row_imgs = [table_image[y:y + h, x:x + w] for (x, y, w, h) in row_sorted]
            cells.append(row_imgs)

        return cells

    def extract_image(self, file: UploadFile) -> str:
        """Extract from an image file using OCR."""
        try:
            file.file.seek(0)
            image = np.array(Image.open(file.file))

            content = self.__ocr_image(image)
            return content

        except Exception as e:
            raise ValueError(f'Failed to extract text from image file: {str(e)}')

    def extract_xlsx(self, file: UploadFile) -> str:
        """Extract raw text from an XLSX file and convert tables to Markdown."""
        try:
            file.file.seek(0)
            workbook = openpyxl.load_workbook(BytesIO(file.file.read()))
            content = []

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]

                # Add sheet name as header
                content.append(f'## Sheet: {sheet_name}')

                # Find the actual data range
                max_row = 0
                max_col = 0
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value is not None:
                            max_row = max(max_row, cell.row)
                            max_col = max(max_col, cell.column)

                if max_row > 0 and max_col > 0:
                    # Extract data as table
                    table_data = []
                    for row in sheet.iter_rows(min_row=1, max_row=max_row, min_col=1, max_col=max_col):
                        row_data = []
                        for cell in row:
                            value = cell.value
                            if value is None:
                                row_data.append('')
                            else:
                                row_data.append(str(value).strip())
                        table_data.append(row_data)

                    # Convert to markdown table
                    if table_data:
                        md_table = self.__table_to_markdown(table_data)
                        if md_table:
                            content.append(md_table)
                else:
                    content.append('(Empty sheet)')

            return '\n\n'.join(content)
        except Exception as e:
            raise ValueError(f'Failed to extract text from XLSX file: {str(e)}')
